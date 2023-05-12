from docx import Document

document = Document()

document.extended_properties.set_property('total_time', '1')
document.extended_properties.set_property('company', 'ACME')
document.extended_properties.set_property('manager', 'Mr. Trance')
document.extended_properties.set_property('application', 'LibreOffice Word')

# Prints some Extended Properties
print("# Pages:", document.extended_properties.pages)
print("# Characters:", document.extended_properties.characters)

# Just sets some of the regular Core Properties
document.core_properties.author = 'John Doe'
document.core_properties.last_modified_by = 'John Doe'

document.add_heading('Document Title', 0)

p = document.add_paragraph('Have a look at the core & extended props of me.')

document.save('core_extended_props_document.docx')
