import docx

document = docx.Document()

# Scenario 1: add a comment (with custom author and initials) on the entire paragraph
paragraph1 = document.add_paragraph('text par 1')
comment = paragraph1.add_comment('comment from scripted author', author='Scripted Author', initials='sa')

# Scenario 2: add a comment (with python-docx as author) only for the run text
paragraph2 = document.add_paragraph('text par 2')
run = paragraph2.add_run('some other')
run.add_comment('comment from py')

document.save('document-comments.docx')
