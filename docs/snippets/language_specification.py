import docx


doc = docx.Document()

# For new document (document-wide):
# Set language value in the documents' default Run's Properties element.
styles_element = doc.styles.element
rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
lang_default = rpr_default.xpath('w:lang')[0]
lang_default.set(docx.oxml.shared.qn('w:val'), 'de-DE')

title = doc.add_paragraph('Rechtschreibprüfung', style='Title')

p1 = doc.add_paragraph(
    'Das ist ein deutscher Satz. '
    'Die Rechtschreibprüfung sollte nichts anstreichen.',
    style='Normal'
    )

# For existing styles:
# For styles without a language value
# you can append one explicitly by
# iterating over those styles in the document.
for my_style in doc.styles:
    style = doc.styles[my_style.name]
    rpr = style.element.get_or_add_rPr()
    lang = docx.oxml.shared.OxmlElement('w:lang')
    if not rpr.xpath('w:lang'):
        lang.set(docx.oxml.shared.qn('w:val'), 'de-DE')
        lang.set(docx.oxml.shared.qn('w:eastAsia'), 'en-US')
        lang.set(docx.oxml.shared.qn('w:bidi'), 'ar-SA')
        rpr.append(lang)

p2 = doc.add_paragraph(
    'This sentence is written in English. '
    'The automatic spell checking should complain, '
    'because all styles’ language was set to German.',
    style='Quote'
    )

# For addressing specifc styles:
# Update (or append to) a specific style,
# e.g. in order to use multiple styles
# to handle more than one language per document.
body_style = doc.styles['Body Text']
body_rpr = body_style.element.get_or_add_rPr()
body_lang = body_rpr.xpath('w:lang')[0]
body_lang.set(docx.oxml.shared.qn('w:val'), 'en-US')

p3 = doc.add_paragraph(
    'This sentence is written again in English. '
    'The automatic spell checking should not complain, '
    'because this style’s language now has been set to English.',
    style='Body Text'
    )

# Run Level:
# For mixing multiple languages
# within the same style per paragraph.
p4 = doc.add_paragraph(style='Body Text')
p4_text = p4.add_run()
p4_text.add_text(
    'On Run Level: This sentence is written once again in English. '
    'Spell check = OK | '
    )
# Add a new run with its language
# differing from the style's language value.
p4_text = p4.add_run()
p4_rpr = p4_text.element.get_or_add_rPr()
p4_run_lang = docx.oxml.shared.OxmlElement('w:lang')
p4_run_lang.set(docx.oxml.shared.qn('w:val'), 'de-DE')
p4_run_lang.set(docx.oxml.shared.qn('w:eastAsia'), 'en-US')
p4_run_lang.set(docx.oxml.shared.qn('w:bidi'), 'ar-SA')
p4_rpr.append(p4_run_lang)
p4_text.add_text(
    'Und das ist noch einmal ein deutscher Satz. '
    'Rechtschreibprüfung = okay'
    )

doc.save('document-different-languages.docx')
